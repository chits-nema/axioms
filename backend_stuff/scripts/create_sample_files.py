from __future__ import annotations

import csv
from pathlib import Path


SAMPLE_DIR = Path(__file__).resolve().parents[1] / "app" / "storage" / "uploads" / "samples"

QUOTE_ROWS = [
    {
        "vendor": "Acme Supplies",
        "product_name": "Industrial Sensor X200",
        "sku": "X200",
        "quantity": 100,
        "unit_price": 42.50,
        "total_price": 4250.00,
        "currency": "EUR",
        "delivery_days": 14,
        "warranty_months": 24,
        "payment_terms": "Net 30",
    },
    {
        "vendor": "Acme Supplies",
        "product_name": "Mounting Kit M7",
        "sku": "M7-KIT",
        "quantity": 100,
        "unit_price": 6.25,
        "total_price": 625.00,
        "currency": "EUR",
        "delivery_days": 14,
        "warranty_months": 12,
        "payment_terms": "Net 30",
    },
]

TEXT_QUOTE = """Vendor: Acme Supplies
Currency: EUR
Quote valid until: 2026-06-30
Payment terms: Net 30
Shipping cost: 150

Item 1:
Product: Industrial Sensor X200
SKU: X200
Quantity: 100
Unit price: 42.50
Total price: 4250.00
Delivery: 14 days
Warranty: 24 months

Item 2:
Product: Mounting Kit M7
SKU: M7-KIT
Quantity: 100
Unit price: 6.25
Total price: 625.00
Delivery: 14 days
Warranty: 12 months
"""


def main() -> None:
    SAMPLE_DIR.mkdir(parents=True, exist_ok=True)

    _create_txt()
    _create_csv()
    _create_xlsx()
    _create_docx()
    _create_pdf()
    _create_png()

    print(f"Created sample files in {SAMPLE_DIR}")


def _create_txt() -> None:
    (SAMPLE_DIR / "sample_quote.txt").write_text(TEXT_QUOTE, encoding="utf-8")


def _create_csv() -> None:
    with (SAMPLE_DIR / "sample_quote.csv").open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(QUOTE_ROWS[0].keys()))
        writer.writeheader()
        writer.writerows(QUOTE_ROWS)


def _create_xlsx() -> None:
    try:
        from openpyxl import Workbook
    except ImportError as exc:
        raise RuntimeError("Install openpyxl first: pip install -r requirements.txt") from exc

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Vendor Quote"
    sheet.append(list(QUOTE_ROWS[0].keys()))
    for row in QUOTE_ROWS:
        sheet.append(list(row.values()))
    workbook.save(SAMPLE_DIR / "sample_quote.xlsx")


def _create_docx() -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Install python-docx first: pip install -r requirements.txt") from exc

    document = Document()
    document.add_heading("Quotation - Acme Supplies", level=1)
    document.add_paragraph("Currency: EUR")
    document.add_paragraph("Quote valid until: 2026-06-30")
    document.add_paragraph("Payment terms: Net 30")

    table = document.add_table(rows=1, cols=len(QUOTE_ROWS[0]))
    for index, header in enumerate(QUOTE_ROWS[0].keys()):
        table.rows[0].cells[index].text = header

    for quote_row in QUOTE_ROWS:
        cells = table.add_row().cells
        for index, value in enumerate(quote_row.values()):
            cells[index].text = str(value)

    document.save(SAMPLE_DIR / "sample_quote.docx")


def _create_pdf() -> None:
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("Install pymupdf first: pip install -r requirements.txt") from exc

    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), TEXT_QUOTE, fontsize=11)
    document.save(SAMPLE_DIR / "sample_quote.pdf")
    document.close()


def _create_png() -> None:
    try:
        from PIL import Image, ImageDraw
    except ImportError as exc:
        raise RuntimeError("Install pillow first: pip install -r requirements.txt") from exc

    image = Image.new("RGB", (1100, 850), "white")
    draw = ImageDraw.Draw(image)
    draw.multiline_text((40, 40), TEXT_QUOTE, fill="black", spacing=8)
    image.save(SAMPLE_DIR / "sample_quote.png")


if __name__ == "__main__":
    main()
